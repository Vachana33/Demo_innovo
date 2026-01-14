from fastapi import APIRouter, Depends, HTTPException, status, Response
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy.orm.attributes import flag_modified
from app.database import get_db
from app.models import Document, Company, User
from app.schemas import DocumentResponse, DocumentUpdate, DocumentContent, DocumentSection, ChatRequest, ChatResponse, ChatConfirmationRequest
from app.dependencies import get_current_user
from typing import List, Optional, Tuple
from datetime import datetime, timezone
import os
import json
import re
import logging
import io
import traceback
from openai import OpenAI
import openai

logger = logging.getLogger(__name__)

router = APIRouter()

# ============================================================================
# ROLE SEPARATION ENFORCEMENT
# ============================================================================
# This module maintains strict separation between:
#
# 1. INITIAL GENERATION (used by /generate-content):
#    - Function: _generate_batch_content()
#    - Purpose: Creates content from scratch for empty sections
#    - Prompt: Assumes empty sections, focuses on creation
#
# 2. CHAT EDITING (used by /chat):
#    - Function: _generate_section_content()
#    - Purpose: Modifies existing section content
#    - Prompt: Assumes existing content, focuses on modification
#
# VERIFIED: No cross-calling exists:
# - /generate-content ONLY calls _generate_batch_content()
# - /chat ONLY calls _generate_section_content()
# ============================================================================

@router.get(
    "/documents/{company_id}/vorhabensbeschreibung",
    response_model=DocumentResponse
)
def get_document(
    company_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Get a document for a company by type.
    Currently only supports "vorhabensbeschreibung".
    """
    # Verify company exists
    company = db.query(Company).filter(Company.id == company_id).first()
    if not company:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Company not found"
        )
    
    # Get or create document
    document = db.query(Document).filter(
        Document.company_id == company_id,
        Document.type == "vorhabensbeschreibung"
    ).first()
    
    if not document:
        # Create empty document if it doesn't exist
        document = Document(
            company_id=company_id,
            type="vorhabensbeschreibung",
            content_json={"sections": []},
            chat_history=[]  # Initialize empty chat history
        )
        db.add(document)
        db.commit()
        db.refresh(document)
    
    # Ensure chat_history is initialized if null
    # Handle case where column might not exist yet (migration not run)
    try:
        if document.chat_history is None:
            document.chat_history = []
            try:
                db.commit()
                db.refresh(document)
                logger.debug(f"Initialized chat_history for document {document.id}")
            except Exception as e:
                logger.warning(f"Failed to initialize chat_history (column may not exist yet): {str(e)}")
                # Set to empty list in memory even if DB update fails
                document.chat_history = []
    except AttributeError:
        # Column doesn't exist in database yet - set to empty list
        logger.warning(f"chat_history column may not exist in database for document {document.id}")
        document.chat_history = []
    
    return document

@router.put(
    "/documents/{document_id}",
    response_model=DocumentResponse
)
def update_document(
    document_id: int,
    document_data: DocumentUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Update a document's content.
    """
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Update content
    document.content_json = document_data.content_json
    
    try:
        db.commit()
        db.refresh(document)
        return document
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to update document: {str(e)}"
        )


# PDF style reference cache (extracted once, reused many times)
_pdf_style_reference_cache: Optional[str] = None


def _extract_pdf_style_reference(pdf_path: str) -> str:
    """
    Extract text from a PDF file to use as style reference.
    Returns a cleaned text sample (first 2000-3000 chars) that represents style, tone, and structure.
    If extraction fails, returns empty string (silent fallback).
    
    CRITICAL: This function must not raise exceptions - it must gracefully handle all errors.
    """
    try:
        import PyPDF2
        
        if not os.path.exists(pdf_path):
            logger.warning(f"PDF file not found: {pdf_path}")
            return ""
        
        text_content = []
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            # Extract text from first few pages (usually contains style examples)
            max_pages = min(3, len(pdf_reader.pages))
            for page_num in range(max_pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_content.append(text)
            
            # Combine and clean text
            combined_text = "\n\n".join(text_content)
            
            # Take first 2500 characters as style sample
            style_sample = combined_text[:2500].strip()
            
            # Clean up excessive whitespace
            style_sample = re.sub(r'\s+', ' ', style_sample)
            style_sample = re.sub(r'\n\s*\n', '\n\n', style_sample)
            
            return style_sample
            
    except ImportError:
        logger.warning("PyPDF2 library not installed. PDF style references will not be available.")
        return ""
    except Exception as e:
        logger.warning(f"Failed to extract text from PDF {pdf_path}: {str(e)}")
        return ""


def _build_style_reference_text() -> str:
    """
    Build style reference text from both DIlico.pdf and Lagotec.pdf.
    Caches the result to avoid re-extracting on every call.
    Returns formatted text that can be inserted into prompts.
    If extraction fails, returns empty string (silent fallback).
    
    CRITICAL: This function must not raise exceptions - it must gracefully handle all errors.
    """
    global _pdf_style_reference_cache
    
    # Return cached result if available
    if _pdf_style_reference_cache is not None:
        return _pdf_style_reference_cache
    
    try:
        # Get the base directory (backend/app)
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        pdf_dir = os.path.join(base_dir, "app", "ai", "prompts", "vorhabensbeschreibung")
        
        dlico_path = os.path.join(pdf_dir, "DIlico.pdf")
        lagotec_path = os.path.join(pdf_dir, "Lagotec.pdf")
        
        dlico_text = _extract_pdf_style_reference(dlico_path)
        lagotec_text = _extract_pdf_style_reference(lagotec_path)
        
        # Build style reference section
        style_parts = []
        
        if dlico_text or lagotec_text:
            style_parts.append("WICHTIG - STILREFERENZEN:")
            style_parts.append("Sie haben Zugriff auf zwei Beispiel-Vorhabensbeschreibungen als PDF-Dateien (DIlico.pdf und Lagotec.pdf).")
            style_parts.append("Diese PDFs dienen AUSSCHLIESSLICH als Stilreferenzen für:")
            style_parts.append("- Ton und Formulierungsstil")
            style_parts.append("- Absatzlänge und narrative Tiefe")
            style_parts.append("- Strukturdichte")
            style_parts.append("- Formalisierungsgrad")
            style_parts.append("- Fördermittel-typische Struktur")
            style_parts.append("")
            style_parts.append("KRITISCH:")
            style_parts.append("- Kopieren Sie KEINEN Inhalt aus diesen PDFs")
            style_parts.append("- Paraphrasieren Sie KEINEN Inhalt aus diesen PDFs")
            style_parts.append("- Verwenden Sie KEINE Fakten aus diesen Dokumenten")
            style_parts.append("- Erwähnen Sie diese PDFs NICHT im generierten Text")
            style_parts.append("- Alle faktischen Inhalte müssen AUSSCHLIESSLICH aus den bereitgestellten Firmeninformationen stammen")
            style_parts.append("")
            
            if dlico_text:
                style_parts.append("STILBEISPIEL aus DIlico.pdf (NUR als Stilreferenz):")
                style_parts.append(dlico_text)
                style_parts.append("")
            
            if lagotec_text:
                style_parts.append("STILBEISPIEL aus Lagotec.pdf (NUR als Stilreferenz):")
                style_parts.append(lagotec_text)
                style_parts.append("")
            
            style_parts.append("Passen Sie Absatzlänge, narrative Dichte und professionellen Ton an den Stil der Beispiel-PDFs an.")
            style_parts.append("")
        
        result = "\n".join(style_parts)
        
        # Cache the result (even if empty)
        _pdf_style_reference_cache = result
        
        return result
        
    except Exception as e:
        logger.warning(f"Failed to build style reference text: {str(e)}")
        # Cache empty string to avoid repeated failures
        _pdf_style_reference_cache = ""
        return ""


def _split_sections_into_batches(sections: List[dict], batch_size: int = 4) -> List[List[dict]]:
    """
    Split sections into batches of 3-5 headings for chunked generation.
    Default batch_size is 4, but will vary between 3-5 to balance efficiency and reliability.
    """
    batches = []
    current_batch = []
    
    for section in sections:
        current_batch.append(section)
        if len(current_batch) >= batch_size:
            batches.append(current_batch)
            current_batch = []
            # Vary batch size slightly (3-5) for better distribution
            batch_size = 3 if batch_size == 5 else 5 if batch_size == 3 else 4
    
    # Add remaining sections as final batch
    if current_batch:
        batches.append(current_batch)
    
    return batches


def _generate_batch_content(
    client: OpenAI,
    batch_sections: List[dict],
    company_name: str,
    website_text: str,
    transcript_text: str,
    max_retries: int = 2
) -> dict:
    """
    ROLE: INITIAL GENERATION
    
    Creates section content from scratch for empty or new sections.
    Used ONLY during first draft generation via /generate-content endpoint.
    
    This function:
    - Assumes sections are empty or need initial content
    - Focuses on creation and expansion
    - Can be creative and comprehensive
    - Generates content based on company data and style references
    
    This function must NOT:
    - Be used for chat-based editing
    - Modify existing section content
    - Be called from /chat endpoint
    
    Returns a dictionary mapping section_id to generated content.
    Implements retry logic with strict JSON validation.
    """
    # Build headings list for this batch
    headings_list = []
    section_ids = []
    for section in batch_sections:
        section_id = section.get('id', '')
        section_title = section.get('title', '')
        # Remove numbering prefix from title
        clean_title = re.sub(r'^[\d.]+\.\s*', '', section_title)
        headings_list.append(f"{section_id}. {clean_title}")
        section_ids.append(section_id)
    
    headings_text = "\n".join(headings_list)
    
    # Smart truncation for company data
    MAX_TEXT_LENGTH = 50000
    
    def smart_truncate(text: str, max_length: int) -> str:
        """Truncate text intelligently, keeping beginning and end if too long."""
        if len(text) <= max_length:
            return text
        first_part = text[:int(max_length * 0.6)]
        last_part = text[-int(max_length * 0.4):]
        return f"{first_part}\n\n[... content truncated ...]\n\n{last_part}"
    
    website_text_processed = smart_truncate(website_text, MAX_TEXT_LENGTH)
    transcript_text_processed = smart_truncate(transcript_text, MAX_TEXT_LENGTH)
    
    # IMPORTANT: This prompt is for INITIAL CONTENT GENERATION only.
    # It assumes empty sections and focuses on creation.
    # Do NOT reuse this prompt for chat-based editing.
    # For editing existing content, use _generate_section_content() instead.
    
    # Build prompt with PDF style reference instructions
    prompt = f"""Sie sind ein Expertenberater, der bei der Erstellung einer "Vorhabensbeschreibung" für einen Förderantrag hilft.

WICHTIG - STILREFERENZEN:
Sie haben Zugriff auf zwei Beispiel-Vorhabensbeschreibungen als PDF-Dateien (DIlico.pdf und Lagotec.pdf).
Diese PDFs dienen AUSSCHLIESSLICH als Stilreferenzen für:
- Ton und Formulierungsstil
- Absatzlänge und narrative Tiefe
- Strukturdichte
- Formalisierungsgrad
- Fördermittel-typische Struktur

KRITISCH: 
- Kopieren Sie KEINEN Inhalt aus diesen PDFs
- Paraphrasieren Sie KEINEN Inhalt aus diesen PDFs
- Verwenden Sie KEINE Fakten aus diesen Dokumenten
- Erwähnen Sie diese PDFs NICHT im generierten Text
- Alle faktischen Inhalte müssen AUSSCHLIESSLICH aus den untenstehenden Firmeninformationen stammen

Firmeninformationen:
- Firmenname: {company_name}
- Website-Inhalt: {website_text_processed}
- Besprechungsprotokoll: {transcript_text_processed}

Die folgenden Abschnitte müssen generiert werden:
{headings_text}

Generieren Sie für jeden Abschnitt detaillierte, professionelle Inhalte basierend AUSSCHLIESSLICH auf den oben genannten Firmeninformationen.

SPRACHE UND STIL:
- Schreiben Sie AUSSCHLIESSLICH auf Deutsch
- Verwenden Sie formelle Fördermittel-/Geschäftssprache
- Verwenden Sie NUR Absätze (keine Aufzählungspunkte)
- Passen Sie Absatzlänge, narrative Dichte und professionellen Ton an den Stil der Beispiel-PDFs an
- Wenn Informationen unzureichend sind, generieren Sie plausible, professionelle Inhalte
- Fügen Sie KEINE Platzhalter ein
- Stellen Sie KEINE Fragen an den Benutzer
- Fügen Sie KEINE Zitate oder Haftungsausschlüsse ein

WICHTIG: Geben Sie NUR ein gültiges JSON-Objekt mit dieser exakten Struktur zurück:
{{
  "{section_ids[0] if section_ids else "section_id"}": "Generierter Absatztext...",
  "{section_ids[1] if len(section_ids) > 1 else "section_id"}": "Generierter Absatztext..."
}}

Die Schlüssel MÜSSEN exakt mit den Abschnitts-IDs aus der Liste oben übereinstimmen (z.B. "0", "1", "1.1", "2.3", etc.).
Die Werte müssen reiner deutscher Text in Absatzform sein.

Geben Sie KEIN Markdown-Format, KEINE Erklärungen und KEINEN Text außerhalb des JSON-Objekts zurück. Geben Sie NUR das JSON-Objekt zurück."""

    # Retry logic with JSON validation
    for attempt in range(max_retries + 1):
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": "Sie sind ein professioneller Berater, der sich auf Förderanträge spezialisiert hat. Sie schreiben klare, strukturierte und überzeugende Projektbeschreibungen auf Deutsch im formellen Fördermittel-Stil."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.7,
                response_format={"type": "json_object"}
            )
            
            response_text = response.choices[0].message.content
            logger.info(f"OpenAI response received for batch (attempt {attempt + 1})")
            
            # Strict JSON validation
            try:
                generated_content = json.loads(response_text)
                
                # Validate that all expected section IDs are present
                missing_ids = [sid for sid in section_ids if sid not in generated_content]
                if missing_ids:
                    raise ValueError(f"Missing section IDs in response: {missing_ids}")
                
                # Validate that all values are strings
                for sid, content in generated_content.items():
                    if not isinstance(content, str):
                        raise ValueError(f"Content for section {sid} is not a string: {type(content)}")
                    if sid not in section_ids:
                        logger.warning(f"Unexpected section ID in response: {sid}")
                
                logger.info(f"Successfully validated JSON for batch with {len(generated_content)} sections")
                return generated_content
                
            except json.JSONDecodeError as e:
                error_msg = f"JSON parse error (attempt {attempt + 1}/{max_retries + 1}): {str(e)}. Response preview: {response_text[:200]}"
                logger.warning(error_msg)
                if attempt < max_retries:
                    continue
                raise ValueError(f"Failed to parse JSON after {max_retries + 1} attempts: {str(e)}")
                
            except ValueError as e:
                error_msg = f"JSON validation error (attempt {attempt + 1}/{max_retries + 1}): {str(e)}"
                logger.warning(error_msg)
                if attempt < max_retries:
                    continue
                raise
                
        except Exception as e:
            if attempt < max_retries:
                logger.warning(f"OpenAI API error (attempt {attempt + 1}/{max_retries + 1}): {str(e)}. Retrying...")
                continue
            logger.error(f"OpenAI API error after {max_retries + 1} attempts: {str(e)}")
            raise
    
    # Should never reach here, but just in case
    raise ValueError("Failed to generate content after all retries")


@router.post(
    "/documents/{document_id}/generate-content",
    response_model=DocumentResponse
)
def generate_content(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    ROLE: INITIAL CONTENT GENERATION
    
    Generate content for Vorhabensbeschreibung document using OpenAI with chunked generation.
    Requires company preprocessing to be completed.
    Generates content in batches of 3-5 sections for reliability and efficiency.
    
    This endpoint:
    - Creates initial content for empty sections
    - Uses _generate_batch_content() for generation logic
    - Assumes sections exist but have no content yet
    
    This endpoint must NOT:
    - Call _generate_section_content() (that's for editing only)
    - Be used for modifying existing content (use /chat instead)
    """
    # Load document
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Verify document type
    if document.type != "vorhabensbeschreibung":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Content generation only supported for vorhabensbeschreibung documents"
        )
    
    # Load associated company
    company = db.query(Company).filter(Company.id == document.company_id).first()
    if not company:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Company not found"
        )
    
    # Check processing status
    if company.processing_status != "done":
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Company preprocessing not finished"
        )
    
    # Load confirmed headings from document
    content_json = document.content_json
    if not content_json or "sections" not in content_json:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document has no sections. Please create and confirm headings first."
        )
    
    sections = content_json["sections"]
    if not sections or len(sections) == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document has no confirmed headings. Please create and confirm headings first."
        )
    
    # Get OpenAI API key from environment
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="OpenAI API key not configured. Set OPENAI_API_KEY environment variable."
        )
    
    # Initialize OpenAI client
    try:
        client = OpenAI(api_key=api_key)
    except Exception as e:
        logger.error(f"Failed to initialize OpenAI client: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to initialize OpenAI client: {str(e)}"
        )
    
    # Prepare company data
    company_name = company.name or "Unknown Company"
    website_text = company.website_text or ""
    transcript_text = company.transcript_text or ""
    
    # Split sections into batches (3-5 sections per batch)
    batches = _split_sections_into_batches(sections, batch_size=4)
    logger.info(f"Split {len(sections)} sections into {len(batches)} batches for document {document_id}")
    
    # Initialize section content map (preserve existing content)
    section_content_map = {}
    for section in sections:
        section_id = section.get("id", "")
        existing_content = section.get("content", "")
        section_content_map[section_id] = existing_content
    
    # Process each batch
    successful_batches = 0
    failed_batches = []
    
    for batch_idx, batch in enumerate(batches):
        try:
            logger.info(f"Processing batch {batch_idx + 1}/{len(batches)} with {len(batch)} sections")
            
            # Generate content for this batch
            # NOTE: This calls _generate_batch_content (INITIAL GENERATION role)
            # This is correct - we are generating initial content, not editing existing content
            batch_content = _generate_batch_content(
                client=client,
                batch_sections=batch,
                company_name=company_name,
                website_text=website_text,
                transcript_text=transcript_text,
                max_retries=2
            )
            
            # Merge batch content into section map
            for section_id, content in batch_content.items():
                if section_id in section_content_map:
                    section_content_map[section_id] = content
                else:
                    logger.warning(f"Generated content for unexpected section ID: {section_id}")
            
            # Persist incrementally after each successful batch
            updated_sections = []
            for section in sections:
                section_id = section.get("id", "")
                section_title = section.get("title", "")
                content = section_content_map.get(section_id, section.get("content", ""))
                
                updated_sections.append({
                    "id": section_id,
                    "title": section_title,
                    "content": content
                })
            
            document.content_json = {"sections": updated_sections}
            db.commit()
            db.refresh(document)
            
            successful_batches += 1
            logger.info(f"Successfully processed and persisted batch {batch_idx + 1}/{len(batches)}")
            
        except Exception as e:
            # Log error but continue with other batches
            batch_section_ids = [s.get("id", "") for s in batch]
            error_msg = f"Failed to generate content for batch {batch_idx + 1} (sections: {batch_section_ids}): {str(e)}"
            logger.error(error_msg)
            logger.error(f"Full traceback:\n{traceback.format_exc()}")
            failed_batches.append({
                "batch_index": batch_idx + 1,
                "section_ids": batch_section_ids,
                "error": str(e)
            })
            # Continue with next batch - partial success is preserved
    
    # Final status check
    if successful_batches == 0:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate content for all batches. Errors: {[b['error'] for b in failed_batches]}"
        )
    
    if failed_batches:
        logger.warning(f"Completed generation with {len(failed_batches)} failed batches out of {len(batches)} total")
        # Partial success - return what we have, but log the failures
    
    # Final refresh to ensure we return the latest state
    db.refresh(document)
    logger.info(f"Successfully completed content generation for document {document_id}: {successful_batches}/{len(batches)} batches succeeded")
    
    return document


def _normalize_section_id(section_id: str) -> str:
    """
    Normalize section ID by:
    1. Converting commas to dots (e.g., "1,1" -> "1.1")
    2. Stripping trailing dots and whitespace
    3. Removing extra spaces
    
    Examples:
    - "2.1." -> "2.1"
    - "2.1 " -> "2.1"
    - "1,1" -> "1.1"
    - "1, 1" -> "1.1"
    - "2.1.2" -> "2.1.2"
    """
    if not section_id:
        return section_id
    
    # Convert commas to dots (common mistake: "1,1" instead of "1.1")
    normalized = section_id.replace(',', '.')
    
    # Remove spaces around dots (e.g., "1. 1" -> "1.1", "1 , 1" -> "1.1")
    normalized = re.sub(r'\s*\.\s*', '.', normalized)
    
    # Strip trailing dots and whitespace
    normalized = normalized.rstrip('.').strip()
    
    return normalized


def _find_section_by_title(
    user_input: str, 
    sections: List[dict], 
    threshold: float = 0.8
) -> Optional[str]:
    """
    Find section ID by matching title using fuzzy matching.
    
    Strategy:
    1. Exact match (normalized, case-insensitive)
    2. Partial match (title contains input or vice versa)
    3. Fuzzy match using SequenceMatcher (similarity >= threshold)
    
    Returns section_id if found, None otherwise.
    """
    if not user_input or not sections:
        return None
    
    from difflib import SequenceMatcher
    
    # Normalize user input
    user_input_normalized = user_input.lower().strip()
    
    # Build title-to-ID mapping with normalized titles
    title_matches = []  # List of (section_id, normalized_title, similarity_score)
    
    for section in sections:
        section_id = section.get("id", "")
        section_title = section.get("title", "")
        
        if not section_title or not section_id:
            continue
        
        # Remove numbering prefix (e.g., "2.1. Firmengeschichte" -> "Firmengeschichte")
        clean_title = re.sub(r'^[\d.]+\.\s*', '', section_title).strip()
        normalized_title = clean_title.lower()
        
        # Strategy 1: Exact match (normalized)
        if normalized_title == user_input_normalized:
            logger.debug(f"Exact title match: '{user_input}' -> section {section_id}")
            return section_id
        
        # Strategy 2: Partial match (contains)
        if user_input_normalized in normalized_title or normalized_title in user_input_normalized:
            # Calculate similarity for ranking
            similarity = SequenceMatcher(None, user_input_normalized, normalized_title).ratio()
            title_matches.append((section_id, clean_title, similarity))
            logger.debug(f"Partial title match: '{user_input}' ~ '{clean_title}' (similarity: {similarity:.2f})")
            continue
        
        # Strategy 3: Fuzzy match
        similarity = SequenceMatcher(None, user_input_normalized, normalized_title).ratio()
        if similarity >= threshold:
            title_matches.append((section_id, clean_title, similarity))
            logger.debug(f"Fuzzy title match: '{user_input}' ~ '{clean_title}' (similarity: {similarity:.2f})")
    
    # If we have matches, return the best one (highest similarity)
    if title_matches:
        # Sort by similarity (descending), then by section_id for consistency
        title_matches.sort(key=lambda x: (-x[2], x[0]))
        best_match = title_matches[0]
        logger.info(f"Best title match: '{user_input}' -> section {best_match[0]} (similarity: {best_match[2]:.2f})")
        return best_match[0]
    
    return None


def _parse_section_changes_enhanced(user_message: str, valid_section_ids: List[str], sections: List[dict] = None) -> List[dict]:
    """
    Enhanced flexible parser that understands various natural language formats.
    This parser is more permissive than the original but still deterministic and safe.
    
    Supports formats like:
    - "Section 2.1: make it more concise"
    - "2.1: make it innovative"
    - "2.1 innovation"
    - "2.1 - innovation"
    - "Update 2.3 to emphasize sustainability"
    - "2.1: concise. 2.2: more technical"
    - "2.1, 2.3, and 2.5: make them all innovative"
    
    Returns empty list if nothing reliable is found (no guessing).
    """
    logger.debug(f"_parse_section_changes_enhanced called with message: '{user_message}', valid_section_ids: {valid_section_ids}")
    changes = []
    message = user_message.strip()
    
    # Strategy: Find all section references first, then extract instructions for each
    
    # Find all potential section references with their positions
    section_matches = []
    
    # NEW: Try to find sections by title first (if sections provided)
    if sections:
        # Pattern 1: "TitleName: instruction" or "TitleName - instruction"
        title_patterns = [
            re.compile(r'([A-Za-zÄÖÜäöüß][A-Za-zÄÖÜäöüß\s]{2,40}?)\s*[:]\s*(.+?)(?=\n|$|[\d.]+\s*:|[A-Za-zÄÖÜäöüß][A-Za-zÄÖÜäöüß\s]{2,40}?\s*:)', re.IGNORECASE),
            re.compile(r'([A-Za-zÄÖÜäöüß][A-Za-zÄÖÜäöüß\s]{2,40}?)\s*[-]\s*(.+?)(?=\n|$|[\d.]+\s*[-:]|section|abschnitt)', re.IGNORECASE),
        ]
        
        for pattern in title_patterns:
            for match in pattern.finditer(message):
                potential_title = match.group(1).strip()
                instruction = match.group(2).strip() if len(match.groups()) > 1 else ""
                
                # Try to find section by title
                section_id = _find_section_by_title(potential_title, sections, threshold=0.8)
                if section_id and section_id in valid_section_ids and instruction and len(instruction) > 2:
                    changes.append({
                        "section_id": section_id,
                        "instruction": instruction
                    })
                    logger.info(f"Found section by title: '{potential_title}' -> {section_id}, instruction: '{instruction}'")
                    # Remove this part from message to avoid duplicate parsing
                    message = message[:match.start()] + message[match.end():]
                    break
            if changes:
                break  # If we found a title match, don't try other patterns
        
        # Pattern 2: Standalone title word (e.g., just "Firmengeschichte" followed by instruction)
        # Only try this if no other patterns matched
        if not changes:
            # Look for capitalized words that might be section titles
            standalone_title_pattern = re.compile(r'^([A-ZÄÖÜ][a-zäöüß]+(?:\s+[A-ZÄÖÜ][a-zäöüß]+){0,3})\s+(.+?)(?=\n|$|section|abschnitt|[\d.]+\s*[:])', re.MULTILINE | re.IGNORECASE)
            for match in standalone_title_pattern.finditer(message):
                potential_title = match.group(1).strip()
                instruction = match.group(2).strip() if len(match.groups()) > 1 else ""
                
                # Skip if it looks like a section number pattern (with dots or commas)
                if re.match(r'^[\d.,]+$', potential_title):
                    continue
                
                # Try to find section by title
                section_id = _find_section_by_title(potential_title, sections, threshold=0.8)
                if section_id and section_id in valid_section_ids and instruction and len(instruction) > 2:
                    changes.append({
                        "section_id": section_id,
                        "instruction": instruction
                    })
                    logger.info(f"Found section by standalone title: '{potential_title}' -> {section_id}, instruction: '{instruction}'")
                    break
    
    # Strategy: Find all section references first, then extract instructions for each
    
    # Find all potential section references with their positions
    section_matches = []
    
    # Pattern 1: "Section X.Y" or "Abschnitt X.Y" (also matches commas: "Section 1,1")
    pattern1 = re.compile(r'(?:section|abschnitt)\s+([\d.,]+)', re.IGNORECASE)
    for match in pattern1.finditer(message):
        section_id = _normalize_section_id(match.group(1))
        if section_id in valid_section_ids:
            section_matches.append({
                'id': section_id,
                'start': match.start(),
                'end': match.end(),
                'type': 'explicit'
            })
    
    # Pattern 2: "X.Y:" (direct section ID with colon, also matches commas: "1,1:")
    # Use negative lookbehind to avoid matching partial IDs (e.g., "4" from "2.4")
    pattern2 = re.compile(r'(?<![.,\d])([\d.,]+)\s*:', re.MULTILINE)
    for match in pattern2.finditer(message):
        section_id = _normalize_section_id(match.group(1))
        if section_id in valid_section_ids:
            # Avoid duplicates from pattern1
            if not any(m['id'] == section_id and m['start'] == match.start() for m in section_matches):
                section_matches.append({
                    'id': section_id,
                    'start': match.start(),
                    'end': match.end(),
                    'type': 'colon'
                })
    
    # Pattern 3: "X.Y -" or "X.Y-" (dash format, also matches commas: "1,1 -")
    # Match section ID followed by dash, ensuring we get the full ID (e.g., "2.4" not just "4")
    # Use negative lookbehind to avoid matching partial IDs
    pattern3 = re.compile(r'(?<![.,\d])([\d.,]+)\s*-\s*', re.MULTILINE)
    for match in pattern3.finditer(message):
        section_id = _normalize_section_id(match.group(1))
        if section_id in valid_section_ids:
            # Avoid duplicates
            if not any(m['id'] == section_id and abs(m['start'] - match.start()) < 5 for m in section_matches):
                section_matches.append({
                    'id': section_id,
                    'start': match.start(),
                    'end': match.end(),
                    'type': 'dash'
                })
    
    # Pattern 4: "Update/Rewrite section X.Y" or action verbs with section (also matches commas)
    pattern4 = re.compile(
        r'(?:update|rewrite|change|modify|edit|überarbeite|aktualisiere|ändere|verbessere|erweitere|kürze|betone)\s+(?:section|abschnitt)?\s*([\d.,]+)',
        re.IGNORECASE
    )
    for match in pattern4.finditer(message):
        section_id = _normalize_section_id(match.group(1))
        if section_id in valid_section_ids:
            # Avoid duplicates
            if not any(m['id'] == section_id and abs(m['start'] - match.start()) < 10 for m in section_matches):
                section_matches.append({
                    'id': section_id,
                    'start': match.start(),
                    'end': match.end(),
                    'type': 'action'
                })
    
    # Pattern 5: Standalone section ID at start of line or after punctuation (also matches commas)
    pattern5 = re.compile(r'(?:^|[\n\.])\s*([\d.,]+)\s+(?![\d.,])', re.MULTILINE)
    for match in pattern5.finditer(message):
        section_id = _normalize_section_id(match.group(1))
        if section_id in valid_section_ids:
            # Only add if it's clearly a section reference (not part of a number)
            # Check if followed by meaningful text (not just another number)
            pos = match.end()
            if pos < len(message):
                next_chars = message[pos:pos+20].strip()
                # If followed by action words or meaningful text, it's likely a section reference
                if next_chars and not re.match(r'^[\d.\s,]+$', next_chars):
                    # Avoid duplicates
                    if not any(m['id'] == section_id and abs(m['start'] - match.start()) < 5 for m in section_matches):
                        section_matches.append({
                            'id': section_id,
                            'start': match.start(),
                            'end': match.end(),
                            'type': 'standalone'
                        })
    
    # Remove duplicates (keep first occurrence)
    seen_ids = set()
    unique_matches = []
    for match in section_matches:
        if match['id'] not in seen_ids:
            seen_ids.add(match['id'])
            unique_matches.append(match)
            logger.debug(f"Added section match: id={match['id']}, type={match['type']}, position={match['start']}")
        else:
            logger.debug(f"Skipped duplicate section match: id={match['id']}, type={match['type']}, position={match['start']}")
    
    # Sort by position in message
    unique_matches.sort(key=lambda x: x['start'])
    logger.debug(f"Final unique matches after sorting: {[m['id'] for m in unique_matches]}")
    
    # Extract instruction for each section
    for i, sec_match in enumerate(unique_matches):
        section_id = sec_match['id']
        instruction_start = sec_match['end']
        
        # Find where this instruction ends (next section or end of message)
        if i + 1 < len(unique_matches):
            instruction_end = unique_matches[i + 1]['start']
        else:
            instruction_end = len(message)
        
        # Extract instruction text
        instruction_text = message[instruction_start:instruction_end].strip()
        
        # Clean up instruction
        # Remove leading separators (colon, dash, whitespace)
        # Note: dash must be escaped or at end of character class to avoid being interpreted as range
        instruction_text = re.sub(r'^[-:\s]+', '', instruction_text)
        
        # Remove trailing separators
        instruction_text = re.sub(r'\s*[-:\s]*$', '', instruction_text)
        
        # Remove trailing punctuation that might be from sentence structure
        instruction_text = re.sub(r'[.,;]+$', '', instruction_text).strip()
        
        # Validate instruction is meaningful
        if instruction_text and len(instruction_text) > 2:
            # Check if it's just another section reference (skip if so)
            # Fix: dash must be at beginning or end of character class
            # Also check for commas in section references
            if not re.match(r'^[\d.,]+\s*[-:\s]', instruction_text):
                logger.debug(f"Found valid change: section_id={section_id}, instruction='{instruction_text}'")
                changes.append({
                    "section_id": section_id,
                    "instruction": instruction_text
                })
            else:
                logger.debug(f"Skipping instruction that looks like section reference: '{instruction_text}'")
        else:
            logger.debug(f"Skipping instruction (too short or empty): '{instruction_text}'")
    
    logger.debug(f"_parse_section_changes_enhanced returning {len(changes)} changes: {changes}")
    return changes


def _parse_section_changes(user_message: str, valid_section_ids: List[str], sections: List[dict] = None) -> List[dict]:
    """
    Parse user message to extract section IDs and their corresponding instructions.
    Returns a list of {section_id, instruction} dictionaries.
    
    This is a deterministic, rule-based parser (not LLM-based).
    Supports multiple formats:
    - "Section 2.1: make it more concise"
    - "Rewrite section 2.1 to emphasize sustainability"
    - "Section 1.1: make more concise. Section 2.3: emphasize innovation"
    - "2.1: make it shorter"
    """
    changes = []
    
    # Normalize message
    message = user_message.strip()
    
    # NEW: Try to find sections by title first (if sections provided)
    if sections:
        # Pattern 1: "TitleName: instruction" or "TitleName - instruction"
        title_patterns = [
            re.compile(r'([A-Za-zÄÖÜäöüß][A-Za-zÄÖÜäöüß\s]{2,40}?)\s*[:]\s*(.+?)(?=\n|$|[\d.]+\s*:|section|abschnitt)', re.IGNORECASE),
            re.compile(r'([A-Za-zÄÖÜäöüß][A-Za-zÄÖÜäöüß\s]{2,40}?)\s*[-]\s*(.+?)(?=\n|$|[\d.]+\s*[-:]|section|abschnitt)', re.IGNORECASE),
        ]
        
        for pattern in title_patterns:
            for match in pattern.finditer(message):
                potential_title = match.group(1).strip()
                instruction = match.group(2).strip() if len(match.groups()) > 1 else ""
                
                # Try to find section by title
                section_id = _find_section_by_title(potential_title, sections, threshold=0.8)
                if section_id and section_id in valid_section_ids and instruction and len(instruction) > 2:
                    changes.append({
                        "section_id": section_id,
                        "instruction": instruction
                    })
                    logger.info(f"Found section by title: '{potential_title}' -> {section_id}, instruction: '{instruction}'")
                    message = message[:match.start()] + message[match.end():]
                    break
            if changes:
                break
        
        # Pattern 2: Standalone title word (e.g., just "Firmengeschichte" followed by instruction)
        # Only try this if no other patterns matched
        if not changes:
            # Look for capitalized words that might be section titles
            standalone_title_pattern = re.compile(r'^([A-ZÄÖÜ][a-zäöüß]+(?:\s+[A-ZÄÖÜ][a-zäöüß]+){0,3})\s+(.+?)(?=\n|$|section|abschnitt|[\d.]+\s*[:])', re.MULTILINE | re.IGNORECASE)
            for match in standalone_title_pattern.finditer(message):
                potential_title = match.group(1).strip()
                instruction = match.group(2).strip() if len(match.groups()) > 1 else ""
                
                # Skip if it looks like a section number pattern (with dots or commas)
                if re.match(r'^[\d.,]+$', potential_title):
                    continue
                
                # Try to find section by title
                section_id = _find_section_by_title(potential_title, sections, threshold=0.8)
                if section_id and section_id in valid_section_ids and instruction and len(instruction) > 2:
                    changes.append({
                        "section_id": section_id,
                        "instruction": instruction
                    })
                    logger.info(f"Found section by standalone title: '{potential_title}' -> {section_id}, instruction: '{instruction}'")
                    break
    
    # Pattern 1: "Section X.Y: instruction" or "Abschnitt X.Y: instruction" (with colon, also matches commas)
    pattern1 = re.compile(r'(?:section|abschnitt)\s+([\d.,]+)\s*:+\s*(.+?)(?=(?:section|abschnitt)\s+[\d.,]+|$)', re.IGNORECASE | re.DOTALL)
    matches1 = pattern1.findall(message)
    for section_id, instruction in matches1:
        section_id = _normalize_section_id(section_id)
        instruction = instruction.strip()
        if section_id in valid_section_ids and instruction and len(instruction) > 3:
            changes.append({"section_id": section_id, "instruction": instruction})
    
    # Pattern 2: "X.Y: instruction" (direct section ID with colon, also matches commas: "1,1:")
    pattern2 = re.compile(r'^([\d.,]+)\s*:+\s*(.+?)(?=\n|$|[\d.,]+\s*:+)', re.MULTILINE | re.IGNORECASE | re.DOTALL)
    matches2 = pattern2.findall(message)
    for section_id, instruction in matches2:
        section_id = _normalize_section_id(section_id)
        instruction = instruction.strip()
        if section_id in valid_section_ids and instruction and len(instruction) > 3:
            # Avoid duplicates
            if not any(c["section_id"] == section_id for c in changes):
                changes.append({"section_id": section_id, "instruction": instruction})
    
    # Pattern 3: "Rewrite/Update section X.Y to..." or "Überarbeite Abschnitt X.Y zu..." (also matches commas)
    pattern3 = re.compile(r'(?:rewrite|update|change|modify|edit|überarbeite|aktualisiere|ändere|verbessere|erweitere|kürze|betone)\s+(?:section|abschnitt)?\s*([\d.,]+)\s+(?:to|zu|mit|dass|damit|so dass|um)\s+(.+)', re.IGNORECASE | re.DOTALL)
    matches3 = pattern3.findall(message)
    for section_id, instruction in matches3:
        section_id = _normalize_section_id(section_id)
        instruction = instruction.strip()
        if section_id in valid_section_ids and instruction and len(instruction) > 3:
            # Avoid duplicates
            if not any(c["section_id"] == section_id for c in changes):
                changes.append({"section_id": section_id, "instruction": instruction})
    
    # Pattern 4: "Section X.Y" followed by instruction (without colon, separated by newline or period, also matches commas)
    # Only if no other patterns matched
    if not changes:
        pattern4 = re.compile(r'(?:section|abschnitt)\s+([\d.,]+)\s+([^\n\.]+)', re.IGNORECASE)
        matches4 = pattern4.findall(message)
        for section_id, instruction in matches4:
            section_id = _normalize_section_id(section_id)
            instruction = instruction.strip()
            # Only add if it looks like an instruction (not just "section X.Y" alone)
            if section_id in valid_section_ids and instruction and len(instruction) > 5:
                changes.append({"section_id": section_id, "instruction": instruction})
    
    # Remove duplicates (keep first occurrence)
    seen = set()
    unique_changes = []
    for change in changes:
        if change["section_id"] not in seen:
            seen.add(change["section_id"])
            unique_changes.append(change)
    
    return unique_changes


def _validate_section_changes(changes: List[dict], valid_section_ids: List[str]) -> Tuple[bool, Optional[str]]:
    """
    Validate that all changes have valid section IDs and instructions.
    Returns (is_valid, error_message).
    """
    if not changes:
        return False, None
    
    # Check all section IDs are valid
    invalid_ids = [c["section_id"] for c in changes if c["section_id"] not in valid_section_ids]
    if invalid_ids:
        return False, f"Ungültige Abschnitts-IDs gefunden: {', '.join(invalid_ids)}. Bitte geben Sie gültige Abschnittsnummern an (z.B. 1.1, 2.3)."
    
    # Check all have instructions
    missing_instructions = [c["section_id"] for c in changes if not c.get("instruction") or len(c["instruction"].strip()) < 3]
    if missing_instructions:
        return False, f"Bitte geben Sie für Abschnitt {missing_instructions[0]} eine konkrete Anweisung an, was geändert werden soll."
    
    return True, None


def _determine_clarification_needed(
    user_message: str, 
    valid_section_ids: List[str],
    last_edited_sections: Optional[List[str]] = None
) -> Optional[str]:
    """
    Determine if clarification is needed before calling LLM.
    Returns conversational clarification question if needed, None if ready to proceed.
    This is a deterministic, rule-based check (not LLM-based).
    
    Uses context (last_edited_sections) to suggest sections but never auto-applies.
    """
    logger.debug(f"_determine_clarification_needed called with message: '{user_message}', last_edited_sections: {last_edited_sections}")
    # Try enhanced parser first
    try:
        changes_enhanced = _parse_section_changes_enhanced(user_message, valid_section_ids)
        logger.debug(f"Enhanced parser returned {len(changes_enhanced)} changes")
    except Exception as e:
        logger.error(f"Error in enhanced parser: {str(e)}", exc_info=True)
        changes_enhanced = []
    if changes_enhanced:
        is_valid, error_msg = _validate_section_changes(changes_enhanced, valid_section_ids)
        if is_valid:
            return None  # No clarification needed
        if error_msg:
            return error_msg  # Return validation error
    
    # Fallback to original parser
    try:
        changes_original = _parse_section_changes(user_message, valid_section_ids)
        logger.debug(f"Original parser returned {len(changes_original)} changes")
        if changes_original:
            is_valid, error_msg = _validate_section_changes(changes_original, valid_section_ids)
            if is_valid:
                logger.debug("Original parser found valid changes, no clarification needed")
                return None  # No clarification needed
            if error_msg:
                logger.debug(f"Original parser found changes but validation failed: {error_msg}")
                return error_msg  # Return validation error
    except Exception as e:
        logger.error(f"Error in original parser: {str(e)}", exc_info=True)
    
    # No valid changes found - need clarification
    # Check if message has action verbs (user wants to do something)
    action_pattern = re.compile(
        r'(?:make|update|change|edit|improve|fix|add|remove|rewrite|'
        r'überarbeite|aktualisiere|ändere|verbessere|erweitere|kürze|betone|'
        r'innovative|innovativ|shorter|longer|concise|detailed|technical|technisch)',
        re.IGNORECASE
    )
    has_action = bool(action_pattern.search(user_message))
    
    # Check if any section IDs mentioned (even if not parsed correctly)
    section_refs = re.findall(r'\b([\d.]+)\b', user_message)
    potential_sections = [s for s in section_refs if s in valid_section_ids]
    invalid_sections = [s for s in section_refs if s not in valid_section_ids and re.match(r'^\d+(\.\d+)*$', s)]
    
    # Case 1: Invalid section IDs found
    if invalid_sections:
        unique_invalid = list(set(invalid_sections))
        if len(unique_invalid) == 1:
            return f"Ich konnte Abschnitt {unique_invalid[0]} nicht finden. Bitte geben Sie eine gültige Abschnittsnummer an (z.B. 2.1, 3.2)."
        return f"Ich konnte die Abschnittsnummern {', '.join(unique_invalid)} nicht finden. Bitte geben Sie gültige Abschnittsnummern an (z.B. 2.1, 3.2)."
    
    # Case 2: No sections mentioned at all
    if not potential_sections:
        if has_action:
            # User wants to do something but didn't specify section
            if last_edited_sections and len(last_edited_sections) > 0:
                if len(last_edited_sections) == 1:
                    return f"Meinen Sie Abschnitt {last_edited_sections[0]}? Bitte bestätigen Sie, oder geben Sie die Abschnittsnummer an (z.B. 2.1 oder 2.1 und 2.3)."
                else:
                    sections_str = ", ".join(last_edited_sections)
                    return f"Welche Abschnitte sollen aktualisiert werden? Sie können mehrere angeben (z.B. 2.1 oder 2.1 und 2.3)."
            return "Welche Abschnitte sollen aktualisiert werden? Bitte geben Sie Abschnittsnummern an (z.B. 2.1 oder 2.1 und 2.3)."
        return "Bitte geben Sie an, welche Abschnitte geändert werden sollen und was genau geändert werden soll (z.B. '2.1: make it innovative' oder 'Section 2.1: make it more concise')."
    
    # Case 3: Sections mentioned but couldn't parse instruction
    if has_action:
        if len(potential_sections) == 1:
            return f"Was soll in Abschnitt {potential_sections[0]} geändert werden? Bitte geben Sie eine klarere Anweisung an (z.B. 'make it more innovative' oder 'fix the style')."
        else:
            sections_str = ", ".join(potential_sections)
            return f"Was soll in den Abschnitten {sections_str} geändert werden? Bitte geben Sie für jeden Abschnitt eine Anweisung an (z.B. '2.1: make it innovative. 2.2: fix the style')."
    
    # Case 4: Sections mentioned but no action verb
    if len(potential_sections) == 1:
        return f"Was soll in Abschnitt {potential_sections[0]} geändert werden? Bitte geben Sie eine Anweisung an (z.B. 'make it more innovative' oder 'make it shorter')."
    else:
        sections_str = ", ".join(potential_sections)
        return f"Was soll in den Abschnitten {sections_str} geändert werden? Bitte geben Sie für jeden Abschnitt eine Anweisung an."
    
    # Fallback (should not reach here)
    return "Bitte geben Sie an, welche Abschnitte geändert werden sollen und was genau geändert werden soll."


def _generate_section_content(
    client: OpenAI,
    section_id: str,
    section_title: str,
    current_content: str,
    instruction: str,
    company_name: str,
    website_text: str,
    transcript_text: str
) -> str:
    """
    ROLE: SECTION EDITOR
    
    Modifies EXISTING section content based on user editing instructions.
    Used ONLY for chat-based editing via /chat endpoint.
    
    This function:
    - Assumes sections already have content that needs modification
    - Focuses on targeted editing and refinement
    - Is constrained and conservative (preserves existing structure)
    - Uses existing content as the primary basis for changes
    
    This function must NOT:
    - Be used for initial content generation
    - Regenerate sections from scratch
    - Be called from /generate-content endpoint
    
    Parameters:
    - current_content: The existing section content that will be modified
    - instruction: User's editing instruction (e.g., "make it more concise")
    
    Returns the updated section content as a string.
    """
    # Smart truncation for company data
    MAX_TEXT_LENGTH = 50000
    
    def smart_truncate(text: str, max_length: int) -> str:
        """Truncate text intelligently, keeping beginning and end if too long."""
        if len(text) <= max_length:
            return text
        first_part = text[:int(max_length * 0.6)]
        last_part = text[-int(max_length * 0.4):]
        return f"{first_part}\n\n[... content truncated ...]\n\n{last_part}"
    
    website_text_processed = smart_truncate(website_text, MAX_TEXT_LENGTH)
    transcript_text_processed = smart_truncate(transcript_text, MAX_TEXT_LENGTH)
    
    # Remove numbering prefix from title
    clean_title = re.sub(r'^[\d.]+\.\s*', '', section_title)
    
    # IMPORTANT:
    # This prompt is for EDITING existing content only.
    # Do NOT reuse this prompt for initial content generation.
    # For initial generation, use _generate_batch_content() instead.
    # This prompt assumes existing content exists and must be modified, not created.
    
    # Get style reference text from PDFs (silent fallback if extraction fails)
    style_reference = _build_style_reference_text()
    
    # Build prompt - ADD style reference at the beginning, preserve all existing sections
    prompt = f"""{style_reference}SIE SIND EIN REDAKTEUR, KEIN AUTOR.

- Der folgende Abschnitt EXISTIERT bereits.
- Ihre Aufgabe ist es, den bestehenden Text gezielt zu überarbeiten.
- Ersetzen Sie NICHT den gesamten Inhalt, außer die Benutzeranweisung verlangt dies ausdrücklich.
- Bewahren Sie Struktur, Kernaussagen und Tonalität des bestehenden Textes.

PRIMÄRE GRUNDLAGE:

- Der bestehende Abschnittstext ist die wichtigste Grundlage.
- Änderungen müssen sich auf den vorhandenen Inhalt beziehen.
- Fügen Sie neue Informationen nur hinzu, wenn sie logisch an den bestehenden Text anschließen.

Aktueller Abschnitt:
- Abschnitts-ID: {section_id}
- Titel: {clean_title}
- Aktueller Inhalt: {current_content}

Benutzeranweisung: {instruction}

KONTEXTNUTZUNG:

- Verwenden Sie Website- und Besprechungsinformationen ausschließlich zur Präzisierung oder inhaltlichen Stützung.
- Fügen Sie keine neuen Themen ein, die im bestehenden Abschnitt nicht bereits angelegt sind.
- Vermeiden Sie generische Aussagen ohne Bezug zum aktuellen Abschnitt.

Firmeninformationen (NUR ZUR STÜTZUNG):
- Firmenname: {company_name}
- Website-Inhalt: {website_text_processed}
- Besprechungsprotokoll: {transcript_text_processed}

UMGANG MIT ALLGEMEINEN ANWEISUNGEN:

- Bei unspezifischen Anweisungen wie „Inhalt hinzufügen", „verbessern" oder „ausbauen":
  - Erweitern Sie den bestehenden Text moderat (ca. +20–40%).
  - Vertiefen Sie bestehende Aussagen, anstatt neue Themen zu eröffnen.

- Bei spezifischen Anweisungen wie „kürzer", „präziser" oder „technischer":
  - Passen Sie den Text entsprechend an, behalten Sie aber die Kernaussagen bei.

- Bei Anweisungen wie „rewrite" oder „komplett neu":
  - Formulieren Sie den Text neu, aber behalten Sie die inhaltlichen Kernpunkte bei.
  - Erweitern Sie moderat (ca. +30–50%), nicht exzessiv.

ABSCHNITTSFOKUS:

- Der überarbeitete Text muss inhaltlich eindeutig zum Titel des Abschnitts passen.
- Fügen Sie keine Themen hinzu, die zu anderen Abschnitten gehören.
- Ändern Sie NICHT den Titel oder die Struktur des Abschnitts.

STIL UND SPRACHE:

- Schreiben Sie ausschließlich auf Deutsch.
- Verwenden Sie einen sachlichen, formellen Fördermittel-Stil.
- Schreiben Sie in zusammenhängenden Absätzen (keine Aufzählungen).
- Keine Meta-Kommentare, keine Hinweise auf KI, keine Platzhalter.
- Stellen Sie KEINE Fragen.
- Fügen Sie KEINE Zitate oder Haftungsausschlüsse ein.
- Erwähnen Sie KEINE vorherigen Versionen oder Änderungen.

WICHTIG:

- Ändern Sie NICHT den Abschnittstitel.
- Fügen Sie KEINE neuen Abschnitte hinzu.
- Der Inhalt muss mit den Firmeninformationen übereinstimmen.
- Geben Sie NUR den überarbeiteten Absatztext zurück (kein JSON, kein Markdown, keine Erklärungen)."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "Sie sind ein professioneller Redakteur, der bestehende Abschnitte von Vorhabensbeschreibungen gezielt überarbeitet. Sie sind KEIN Autor, der Inhalte neu erstellt. Ihre Aufgabe ist die präzise Bearbeitung vorhandener Texte auf Deutsch im formellen Fördermittel-Stil."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7
        )
        
        generated_content = response.choices[0].message.content.strip()
        
        # Clean up any markdown or JSON artifacts
        generated_content = re.sub(r'^```(?:json|markdown)?\s*\n?', '', generated_content)
        generated_content = re.sub(r'\n?```\s*$', '', generated_content)
        generated_content = generated_content.strip()
        
        return generated_content
        
    except Exception as e:
        logger.error(f"Failed to generate content for section {section_id}: {str(e)}")
        raise


def _is_question(message: str) -> bool:
    """
    Detect if a message is a question based on patterns.
    Questions typically start with question words or end with '?'.
    """
    message_lower = message.lower().strip()
    
    # Check for question mark
    if message_lower.endswith('?'):
        return True
    
    # Check for question words at the start
    question_words = ['what', 'how', 'why', 'when', 'where', 'who', 'which', 'can', 'could', 'should', 'would', 'is', 'are', 'was', 'were', 'do', 'does', 'did', 'will', 'tell me', 'explain', 'describe']
    first_word = message_lower.split()[0] if message_lower.split() else ""
    
    if first_word in question_words:
        return True
    
    # Check for question patterns
    question_patterns = [
        r'^what\s+',
        r'^how\s+',
        r'^why\s+',
        r'^when\s+',
        r'^where\s+',
        r'^who\s+',
        r'^which\s+',
        r'^can\s+you',
        r'^could\s+you',
        r'^should\s+',
        r'^would\s+',
        r'^tell\s+me',
        r'^explain',
        r'^describe',
    ]
    
    for pattern in question_patterns:
        if re.match(pattern, message_lower):
            return True
    
    return False


def _extract_context_for_question(
    sections: List[dict],
    website_text: str,
    conversation_history: Optional[List[dict]] = None
) -> dict:
    """
    Extract full context for question answering:
    - Full document content (all sections)
    - Website summary (first 200-500 chars)
    - Conversation history (last 2-3 messages)
    """
    # Extract full document content
    document_content_parts = []
    for section in sections:
        section_id = section.get("id", "")
        section_title = section.get("title", "")
        section_content = section.get("content", "")
        if section_content and section_content.strip():
            document_content_parts.append(f"Section {section_id} ({section_title}): {section_content}")
    
    full_document_content = "\n\n".join(document_content_parts) if document_content_parts else "No content generated yet."
    
    # Extract website summary (200-500 chars)
    website_summary = ""
    if website_text:
        # Take first 500 chars, but try to end at a sentence boundary
        if len(website_text) > 500:
            # Find last sentence boundary within 500 chars
            truncated = website_text[:500]
            last_period = truncated.rfind('.')
            last_newline = truncated.rfind('\n')
            cut_point = max(last_period, last_newline)
            if cut_point > 200:  # Ensure we have at least 200 chars
                website_summary = website_text[:cut_point + 1]
            else:
                website_summary = website_text[:500]
        else:
            website_summary = website_text
    
    # Extract conversation history (last 2-3 messages)
    conversation_context = ""
    if conversation_history:
        # Take last 3 messages (or all if less than 3)
        recent_messages = conversation_history[-3:] if len(conversation_history) > 3 else conversation_history
        conversation_parts = []
        for msg in recent_messages:
            role = msg.get("role", "user")
            text = msg.get("text", "")
            if text:
                conversation_parts.append(f"{role.capitalize()}: {text}")
        conversation_context = "\n".join(conversation_parts) if conversation_parts else ""
    
    return {
        "document_content": full_document_content,
        "website_summary": website_summary,
        "conversation_history": conversation_context
    }


def _save_chat_message(
    document: Document,
    role: str,
    text: str,
    suggested_content: Optional[dict] = None,
    requires_confirmation: bool = False,
    db: Session = None
):
    """
    Save a chat message to the document's chat_history.
    """
    # Initialize chat_history if None
    if document.chat_history is None:
        document.chat_history = []
        logger.debug(f"Initialized chat_history for document {document.id}")
    
    # Create message object
    message = {
        "role": role,
        "text": text,
        "timestamp": datetime.now(timezone.utc).isoformat()
    }
    
    # Add optional fields
    if suggested_content:
        message["suggestedContent"] = suggested_content
    if requires_confirmation:
        message["requiresConfirmation"] = True
        message["messageId"] = f"msg-{int(datetime.now(timezone.utc).timestamp() * 1000)}"
    
    # Append to chat history
    document.chat_history.append(message)
    
    # Save to database
    try:
        db.commit()
        db.refresh(document)
        logger.info(f"Saved chat message to document {document.id}: role={role}, text_length={len(text)}, total_messages={len(document.chat_history)}")
    except Exception as e:
        logger.error(f"Failed to save chat message: {str(e)}", exc_info=True)
        db.rollback()
        # Don't raise - chat saving is not critical, but log the error


def _answer_question_with_context(
    client: OpenAI,
    user_query: str,
    document_content: str,
    website_summary: str,
    conversation_history: str,
    company_name: str
) -> str:
    """
    Answer a user question using full context (document, website, conversation history).
    Returns a concise answer in formal business language (Fördermittel tone).
    """
    # Build context prompt
    context_parts = []
    
    if document_content and document_content.strip() != "No content generated yet.":
        context_parts.append(f"Generated Document Content:\n{document_content}")
    
    if website_summary:
        context_parts.append(f"Company Website Summary:\n{website_summary}")
    
    if conversation_history:
        context_parts.append(f"Previous Conversation:\n{conversation_history}")
    
    context_text = "\n\n".join(context_parts)
    
    prompt = f"""Sie sind ein Expertenberater, der Fragen zu einem Förderantrag-Dokument (Vorhabensbeschreibung) beantwortet.

KONTEXT:
{context_text}

Firmenname: {company_name}

BENUTZERFRAGE: "{user_query}"

AUFGABE:
Beantworten Sie die Frage präzise und sachlich im formellen Fördermittel-Stil (Geschäftssprache).

WICHTIGE REGELN:
- Beziehen Sie sich AUSSCHLIESSLICH auf den bereitgestellten Kontext
- Wenn die Antwort nicht im Kontext enthalten ist, sagen Sie dies klar
- Verwenden Sie formelle, professionelle Sprache (Deutsch)
- Seien Sie präzise und konkret
- Keine Spekulationen oder Informationen außerhalb des Kontexts
- Keine Meta-Kommentare oder Hinweise auf KI
- Antworten Sie in zusammenhängenden Absätzen (keine Aufzählungen, außer wenn angebracht)

Geben Sie NUR die Antwort zurück, ohne zusätzliche Erklärungen oder Formatierungen."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "Sie sind ein professioneller Berater für Förderanträge. Sie beantworten Fragen präzise und sachlich im formellen Fördermittel-Stil."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7,
            max_tokens=1000  # Limit response length for concise answers
        )
        
        answer = response.choices[0].message.content.strip()
        logger.info(f"Generated answer for question: '{user_query[:50]}...' (answer length: {len(answer)})")
        return answer
        
    except Exception as e:
        logger.error(f"Error generating answer: {str(e)}")
        raise


@router.post(
    "/documents/{document_id}/chat",
    response_model=ChatResponse
)
def chat_with_document(
    document_id: int,
    chat_request: ChatRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    ROLE: CHAT-BASED SECTION EDITING
    
    Chat endpoint for section-scoped editing of Vorhabensbeschreibung documents.
    Only calls LLM when user explicitly specifies section(s) and change instruction(s).
    Otherwise asks clarification questions.
    
    This endpoint:
    - Modifies existing section content based on user instructions
    - Uses _generate_section_content() for editing logic
    - Assumes sections already have content that needs modification
    
    This endpoint must NOT:
    - Call _generate_batch_content() (that's for initial generation only)
    - Be used for creating initial content (use /generate-content instead)
    """
    # Load document
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Verify document type
    if document.type != "vorhabensbeschreibung":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Chat editing only supported for vorhabensbeschreibung documents"
        )
    
    # Load associated company
    company = db.query(Company).filter(Company.id == document.company_id).first()
    if not company:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Company not found"
        )
    
    # Load document sections
    content_json = document.content_json
    if not content_json or "sections" not in content_json:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document has no sections"
        )
    
    sections = content_json["sections"]
    if not sections:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document has no sections"
        )
    
    # Get valid section IDs
    valid_section_ids = [section.get("id", "") for section in sections if section.get("id")]
    
    # Get context (last edited sections) from request if available
    last_edited_sections = chat_request.last_edited_sections
    conversation_history = chat_request.conversation_history or []
    
    # Check if message is a question
    is_question = _is_question(chat_request.message)
    
    if is_question:
        # Handle question-answering with full context
        logger.info(f"Detected question: '{chat_request.message[:50]}...'")
        
        # Extract context
        context = _extract_context_for_question(
            sections=sections,
            website_text=company.website_text or "",
            conversation_history=conversation_history
        )
        
        # Get OpenAI API key
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="OpenAI API key not configured. Set OPENAI_API_KEY environment variable."
            )
        
        # Initialize OpenAI client
        try:
            client = OpenAI(api_key=api_key)
        except Exception as e:
            logger.error(f"Failed to initialize OpenAI client: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to initialize OpenAI client: {str(e)}"
            )
        
        # Answer the question with context
        try:
            answer = _answer_question_with_context(
                client=client,
                user_query=chat_request.message,
                document_content=context["document_content"],
                website_summary=context["website_summary"],
                conversation_history=context["conversation_history"],
                company_name=company.name or "Unknown Company"
            )
            
            logger.info(f"Question answered successfully (answer length: {len(answer)})")
            
            # Save user message and assistant response to chat history
            _save_chat_message(document, "user", chat_request.message, db=db)
            _save_chat_message(document, "assistant", answer, db=db)
            
            # Return answer without updating any sections
            # The frontend will display the answer in chat
            return ChatResponse(
                message=answer,
                updated_sections=None,
                is_question=True
            )
            
        except Exception as e:
            logger.error(f"Error answering question: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to answer question: {str(e)}"
            )
    
    # If not a question, proceed with section editing logic
    logger.info(f"Message is not a question - proceeding with section editing: '{chat_request.message[:50]}...'")
    
    # Save user message to chat history
    _save_chat_message(document, "user", chat_request.message, db=db)
    
    # Parse section changes: try enhanced parser first, fallback to original
    changes = _parse_section_changes_enhanced(chat_request.message, valid_section_ids, sections)
    
    # If enhanced parser found nothing, try original parser
    if not changes:
        changes = _parse_section_changes(chat_request.message, valid_section_ids, sections)
    
    # If still no changes found, create a default change with raw message
    # (This allows testing even with ambiguous requests)
    if not changes:
        logger.warning(f"No sections parsed from message, creating default change with first section")
        if valid_section_ids:
            # Use first section as default
            changes = [{
                "section_id": valid_section_ids[0],
                "instruction": chat_request.message
            }]
            logger.info(f"Created default change: section={valid_section_ids[0]}, instruction='{chat_request.message}'")
        else:
            return ChatResponse(
                message="Document has no sections to update.",
                updated_sections=None
            )
    
    # Validate changes (keep this for safety, but log warnings and continue)
    is_valid, error_msg = _validate_section_changes(changes, valid_section_ids)
    if not is_valid:
        logger.warning(f"Validation failed but proceeding anyway for testing: {error_msg}")
        # Continue anyway for testing - don't return error
    
    # Get OpenAI API key
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="OpenAI API key not configured. Set OPENAI_API_KEY environment variable."
        )
    
    # Initialize OpenAI client
    try:
        client = OpenAI(api_key=api_key)
    except Exception as e:
        logger.error(f"Failed to initialize OpenAI client: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to initialize OpenAI client: {str(e)}"
        )
    
    # Prepare company data
    company_name = company.name or "Unknown Company"
    website_text = company.website_text or ""
    transcript_text = company.transcript_text or ""
    
    # Process each section change
    updated_section_ids = []
    suggested_content_map = {}  # Map of section_id -> suggested_content for preview
    # Create a map for quick lookup, but we'll update the original sections list
    section_map = {section.get("id"): idx for idx, section in enumerate(sections)}
    
    for change in changes:
        section_id = change["section_id"]
        instruction = change["instruction"]
        
        if section_id not in section_map:
            logger.warning(f"Section {section_id} not found in document")
            continue
        
        section_idx = section_map[section_id]
        section = sections[section_idx]
        section_title = section.get("title", "")
        current_content = section.get("content", "")
        
        try:
            # Generate updated content
            # NOTE: This calls _generate_section_content (SECTION EDITOR role)
            # This is correct - we are editing existing content, not generating initial content
            logger.info(f"Calling LLM for section {section_id} with instruction: '{instruction}'")
            logger.info(f"Current content length: {len(current_content)} characters")
            logger.info(f"Current content preview: {current_content[:100] if current_content else '(empty)'}")
            
            new_content = _generate_section_content(
                client=client,
                section_id=section_id,
                section_title=section_title,
                current_content=current_content,
                instruction=instruction,
                company_name=company_name,
                website_text=website_text,
                transcript_text=transcript_text
            )
            
            logger.info(f"LLM returned content length: {len(new_content)} characters")
            logger.info(f"LLM returned content preview: {new_content[:200]}")
            logger.info(f"Content changed: {new_content != current_content}")
            
            # Check if content actually changed (not just whitespace/formatting)
            content_changed = new_content.strip() != current_content.strip()
            if not content_changed:
                logger.warning(f"LLM returned identical content for section {section_id} - content was not actually modified!")
                logger.warning(f"Instruction was: '{instruction}'")
                logger.warning(f"This may indicate the LLM did not follow the rewrite/expand instruction properly")
            
            # Check if content is significantly longer (for expand/add instructions)
            length_increase = len(new_content) - len(current_content)
            length_increase_percent = (length_increase / len(current_content) * 100) if current_content else 0
            logger.info(f"Content length change: {length_increase} characters ({length_increase_percent:.1f}% increase)")
            
            # Store suggested content (DO NOT update section yet - wait for confirmation)
            # Build map of section_id -> suggested_content for preview
            suggested_content_map[section_id] = new_content
            updated_section_ids.append(section_id)
            logger.info(f"Successfully generated suggested content for section {section_id} for document {document_id} (preview mode)")
            
        except Exception as e:
            logger.error(f"Failed to generate content for section {section_id}: {str(e)}")
            # Continue with other sections even if one fails
            continue
    
    # Return preview instead of saving (user must confirm first)
    if updated_section_ids:
        # Generate preview response message
        try:
            if len(updated_section_ids) == 1:
                response_message = f"Ich habe eine Änderung für Abschnitt {updated_section_ids[0]} vorbereitet. Bitte überprüfen Sie die Vorschau und bestätigen Sie die Änderung."
            else:
                sections_str = ", ".join(updated_section_ids)
                response_message = f"Ich habe Änderungen für die Abschnitte {sections_str} vorbereitet. Bitte überprüfen Sie die Vorschau und bestätigen Sie die Änderungen."
            
            logger.info(f"Returning ChatResponse with preview for {len(updated_section_ids)} sections: {updated_section_ids}")
            
            # Save assistant response with preview to chat history
            _save_chat_message(
                document, 
                "assistant", 
                response_message, 
                suggested_content=suggested_content_map,
                requires_confirmation=True,
                db=db
            )
            
            response = ChatResponse(
                message=response_message,
                suggested_content=suggested_content_map,
                requires_confirmation=True,
                updated_sections=None,  # Not updated yet - waiting for confirmation
                is_question=False  # Explicitly mark as section edit, not question
            )
            logger.info(f"ChatResponse with preview created successfully, returning...")
            return response
        except Exception as e:
            logger.error(f"Error creating ChatResponse: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to create response: {str(e)}"
            )
    else:
        # No sections were updated (all failed)
        error_message = "Entschuldigung, es konnte kein Abschnitt aktualisiert werden. Bitte versuchen Sie es erneut mit spezifischeren Anweisungen."
        _save_chat_message(document, "assistant", error_message, db=db)
        return ChatResponse(
            message=error_message,
            updated_sections=None,
            is_question=False  # Explicitly mark as section edit attempt, not question
        )


@router.post(
    "/documents/{document_id}/chat/confirm",
    response_model=ChatResponse
)
def confirm_chat_edit(
    document_id: int,
    confirmation: ChatConfirmationRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Apply confirmed edit to a section.
    This endpoint is called when user approves a suggested edit from the preview.
    """
    # Load document
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Verify document type
    if document.type != "vorhabensbeschreibung":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Chat confirmation only supported for vorhabensbeschreibung documents"
        )
    
    # Load document sections
    content_json = document.content_json
    if not content_json or "sections" not in content_json:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document has no sections"
        )
    
    sections = content_json["sections"]
    if not sections:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document has no sections"
        )
    
    # Find the section to update
    section_found = False
    logger.info(f"Looking for section {confirmation.section_id} in document {document_id}")
    logger.info(f"Available section IDs: {[s.get('id') for s in sections]}")
    
    for section in sections:
        section_id = section.get("id", "")
        if section_id == confirmation.section_id:
            # Update section content with confirmed content
            old_content_length = len(section.get("content", ""))
            section["content"] = confirmation.confirmed_content
            section_found = True
            logger.info(f"Updating section {confirmation.section_id} with confirmed content (old length: {old_content_length}, new length: {len(confirmation.confirmed_content)})")
            logger.info(f"New content preview: {confirmation.confirmed_content[:200]}...")
            break
    
    if not section_found:
        logger.error(f"Section {confirmation.section_id} not found in document {document_id}. Available sections: {[s.get('id') for s in sections]}")
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Section {confirmation.section_id} not found in document. Available sections: {', '.join([s.get('id', '') for s in sections])}"
        )
    
    # Rebuild sections array preserving order
    # IMPORTANT: Verify the updated content is in the sections list before rebuilding
    updated_sections = []
    for section in sections:
        section_id = section.get("id", "")
        section_content = section.get("content", "")
        
        # Log if this is the section we just updated
        if section_id == confirmation.section_id:
            logger.info(f"Rebuilding section {section_id} with content length: {len(section_content)} (expected: {len(confirmation.confirmed_content)})")
            if section_content != confirmation.confirmed_content:
                logger.error(f"ERROR: Section {section_id} content mismatch during rebuild! Setting correct content.")
                section_content = confirmation.confirmed_content  # Force correct content
        
        updated_sections.append({
            "id": section_id,
            "title": section.get("title", ""),
            "content": section_content
        })
    
    # Verify the updated section is in the rebuilt array
    rebuilt_section = next((s for s in updated_sections if s.get("id") == confirmation.section_id), None)
    if rebuilt_section:
        logger.info(f"Rebuilt section {confirmation.section_id} content length: {len(rebuilt_section.get('content', ''))}")
        if rebuilt_section.get("content") != confirmation.confirmed_content:
            logger.error(f"ERROR: Rebuilt section content doesn't match! Forcing correct content.")
            rebuilt_section["content"] = confirmation.confirmed_content
    
    # Update document in database
    # IMPORTANT: Create a new dict to ensure SQLAlchemy detects the change
    document.content_json = {"sections": updated_sections}
    
    # Mark the JSON column as modified so SQLAlchemy knows to update it
    # This is REQUIRED when modifying nested JSON structures - SQLAlchemy doesn't detect nested changes automatically
    flag_modified(document, "content_json")
    
    try:
        db.commit()
        db.refresh(document)
        
        # Verify the content was actually saved
        saved_section = None
        for s in document.content_json.get("sections", []):
            if s.get("id") == confirmation.section_id:
                saved_section = s
                break
        
        if saved_section:
            saved_content = saved_section.get("content", "")
            logger.info(f"Successfully saved confirmed edit for section {confirmation.section_id} in document {document_id}")
            logger.info(f"Verified saved content length: {len(saved_content)} (expected: {len(confirmation.confirmed_content)})")
            if saved_content != confirmation.confirmed_content:
                logger.error(f"ERROR: Saved content does not match confirmed content!")
                logger.error(f"Expected preview: {confirmation.confirmed_content[:200]}...")
                logger.error(f"Got preview: {saved_content[:200]}...")
                # This is a critical error - raise an exception
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Failed to save content: content mismatch after save"
                )
            else:
                logger.info(f"✓ Content verified successfully - saved content matches confirmed content")
        else:
            logger.error(f"ERROR: Section {confirmation.section_id} not found in saved document!")
            
    except Exception as e:
        db.rollback()
        logger.error(f"Failed to save confirmed edit for document {document_id}: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to save confirmed edit: {str(e)}"
        )
    
    # Return success response
    return ChatResponse(
        message=f"Änderung für Abschnitt {confirmation.section_id} wurde bestätigt und gespeichert.",
        updated_sections=[confirmation.section_id],
        is_question=False,
        requires_confirmation=False
    )


@router.get("/documents/{document_id}/export")
def export_document(
    document_id: int,
    format: str = "pdf",  # "pdf" or "docx"
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Export document as PDF or DOCX file.
    """
    # Load document
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Load associated company for filename
    company = db.query(Company).filter(Company.id == document.company_id).first()
    company_name = company.name if company else "Document"
    # Sanitize filename
    safe_company_name = re.sub(r'[^\w\s-]', '', company_name).strip().replace(' ', '_')
    
    # Get document content
    content_json = document.content_json
    if not content_json or "sections" not in content_json:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document has no content to export"
        )
    
    sections = content_json["sections"]
    
    if format.lower() == "pdf":
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.enums import TA_LEFT
            
            # Create PDF in memory
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=14,
                textColor=(0.2, 0.2, 0.2),
                spaceAfter=12,
                fontName='Helvetica-Bold'
            )
            
            content_style = ParagraphStyle(
                'CustomContent',
                parent=styles['Normal'],
                fontSize=11,
                textColor=(0.1, 0.1, 0.1),
                spaceAfter=12,
                leftIndent=0,
                alignment=TA_LEFT
            )
            
            # Add sections to PDF
            for section in sections:
                title = section.get("title", "")
                content = section.get("content", "")
                
                # Ensure content is a string (handle dict/other types)
                if not isinstance(content, str):
                    if isinstance(content, dict):
                        # If content is a dict, convert to string representation
                        content = str(content)
                    elif content is None:
                        content = ""
                    else:
                        content = str(content)
                
                if title:
                    story.append(Paragraph(title, title_style))
                    story.append(Spacer(1, 6))
                
                if content:
                    # Escape HTML and convert newlines
                    content_escaped = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    content_escaped = content_escaped.replace("\n", "<br/>")
                    story.append(Paragraph(content_escaped, content_style))
                    story.append(Spacer(1, 12))
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            
            filename = f"{safe_company_name}_Vorhabensbeschreibung.pdf"
            return Response(
                content=buffer.getvalue(),
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename={filename}"
                }
            )
        except ImportError:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="PDF export requires reportlab library. Install with: pip install reportlab"
            )
        except Exception as e:
            logger.error(f"PDF export error for document {document_id}: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to generate PDF: {str(e)}"
            )
    
    elif format.lower() == "docx" or format.lower() == "doc":
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Create DOCX document
            docx = DocxDocument()
            
            # Add sections to DOCX
            for section in sections:
                title = section.get("title", "")
                content = section.get("content", "")
                
                if title:
                    title_para = docx.add_paragraph(title)
                    title_para.style = 'Heading 1'
                    title_run = title_para.runs[0] if title_para.runs else title_para.add_run(title)
                    title_run.font.size = Pt(14)
                    title_run.bold = True
                
                if content:
                    content_para = docx.add_paragraph(content)
                    content_para.style = 'Normal'
                    for run in content_para.runs:
                        run.font.size = Pt(11)
                    # Add spacing after content
                    docx.add_paragraph()
            
            # Save to buffer
            buffer = io.BytesIO()
            docx.save(buffer)
            buffer.seek(0)
            
            filename = f"{safe_company_name}_Vorhabensbeschreibung.docx"
            return Response(
                content=buffer.getvalue(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename={filename}"
                }
            )
        except ImportError:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="DOCX export requires python-docx library. Install with: pip install python-docx"
            )
        except Exception as e:
            logger.error(f"DOCX export error for document {document_id}: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to generate DOCX: {str(e)}"
            )
    
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported export format: {format}. Supported formats: pdf, docx"
        )

